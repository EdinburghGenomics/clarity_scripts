#!/usr/bin/env python
import barcode

from barcode.writer import ImageWriter

from docx import Document
from docx.shared import Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from EPPs.common import SendMailEPP


class GenerateTrackingLetter (SendMailEPP):
    #Automatically generates the tracking letter sent to customers for tube shipments by populating a word template
    #with a 128 format barcode containing the container name

    # additional argument required to obtain the file location for newly created tracking letter in the LIMS step
    def __init__(self, argv=None):
        super().__init__(argv)
        self.letter = self.cmd_args.letter

    @staticmethod
    def add_args(argparser):
        argparser.add_argument(
            '-t', '--letter', type=str, required=True, help='Tracking letter generated by the LIMS'
        )

    def _run(self):

        # obtain all of the inputs for the step
        all_inputs = self.process.all_inputs(unique=True)


        # check only one container (rack) is present in the step
        container_names=set()

        for artifact in all_inputs:
            container_names.add(artifact.container.name)

        if len(container_names) >1:
            raise ValueError('Only 1 rack is permitted. Multiple racks are present')


        #generate barcode to be inserted into template
        barcode.PROVIDED_BARCODES
        [ u'code128']

        EAN = barcode.get_barcode_class('code128')

        ean = EAN(list(container_names)[0], writer=ImageWriter())
        save_options = {'font_size':20,
                      'text_distance':2,
                      'module_height':15,
                      'module_width':0.3}
        ean.save('code128', options=save_options)

        document= Document(self.get_config(config_heading_1='file_templates', config_heading_2='tracking_letter'))

        for paragraph in document.paragraphs:
            if 'The barcode(s) above provides confirmation' in paragraph.text:
                p=paragraph.insert_paragraph_before('')
                p=p.insert_paragraph_before('')
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                r=p.add_run()
                r.add_picture('code128.png',width=Cm(5))

        document.save(self.letter+'_Edinburgh_Genomics_Sample_Tracking_Letter_'+all_inputs[0].samples[0].project.name+'.docx')

if __name__ == '__main__':
    GenerateTrackingLetter().run()